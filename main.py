"""
YCM Researcher Enhanced - Main Entry Point
整合 gpt-researcher-clean 的功能与 YCM 的前端界面
"""

import os
import sys
import logging
from pathlib import Path
from dotenv import load_dotenv
from typing import Union, Dict, Any
import json

# 添加父目录到 Python 路径，以便导入 ycm_researcher 模块
parent_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.append(parent_dir)

# 创建日志目录
logs_dir = Path("logs")
logs_dir.mkdir(exist_ok=True)

# 创建缓存目录
cache_dir = Path("cache/research")
cache_dir.mkdir(parents=True, exist_ok=True)

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler("logs/app.log"), logging.StreamHandler()],
)

# 抑制过于详细的日志
logging.getLogger("fontTools").setLevel(logging.WARNING)
logging.getLogger("fontTools.subset").setLevel(logging.WARNING)
logging.getLogger("fontTools.ttLib").setLevel(logging.WARNING)

# 创建日志实例
logger = logging.getLogger(__name__)

# 加载环境变量
load_dotenv()

try:
    # 导入必要的库
    from fastapi import (
        FastAPI,
        Request,
        WebSocket,
        WebSocketDisconnect,
        Body,
        Response,
        HTTPException,
        Depends,
        BackgroundTasks,
        File,
        UploadFile,
        Form,
    )
    from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
    from fastapi.staticfiles import StaticFiles
    from fastapi.templating import Jinja2Templates
    from fastapi.middleware.cors import CORSMiddleware
    import uvicorn
    import time
    import uuid
    import asyncio
    from typing import List, Dict, Any, Optional, Union
    from pydantic import BaseModel
    import aiohttp
    import markdown
    import re
    import shutil
    from datetime import datetime
    import requests
    from urllib.parse import quote

    # 导入 ycm_researcher 模块
    from ycm_researcher.agent import YCMResearcher
    from ycm_researcher.utils.enum import ReportType, ReportSource, Tone
    from ycm_researcher.config.config import Config

    # WebSocket 管理器
    class WebSocketManager:
        def __init__(self):
            self.active_connections = []

        async def connect(self, websocket: WebSocket):
            await websocket.accept()
            self.active_connections.append(websocket)

        def disconnect(self, websocket: WebSocket):
            self.active_connections.remove(websocket)

        async def broadcast(self, message: dict):
            for connection in self.active_connections:
                await connection.send_json(message)

    # 研究进度处理器
    class ResearchProgressHandler:
        def __init__(self, websocket_manager):
            self.websocket_manager = websocket_manager

        async def on_tool_start(self, tool_name, **kwargs):
            await self.websocket_manager.broadcast(
                {
                    "type": "progress",
                    "data": {"message": f"使用工具: {tool_name}", "details": kwargs},
                }
            )

        async def on_agent_action(self, action, **kwargs):
            await self.websocket_manager.broadcast(
                {
                    "type": "progress",
                    "data": {"message": f"执行动作: {action}", "details": kwargs},
                }
            )

        async def on_research_step(self, step, details):
            await self.websocket_manager.broadcast(
                {
                    "type": "progress",
                    "data": {"message": f"研究步骤: {step}", "details": details},
                }
            )

        async def send_progress(self, message):
            await self.websocket_manager.broadcast(
                {"type": "progress", "data": {"message": message}}
            )

    # 创建 WebSocket 管理器实例
    manager = WebSocketManager()
    progress_handler = ResearchProgressHandler(manager)

    # 创建 FastAPI 应用
    app = FastAPI(title="YCM Researcher Enhanced")

    # 添加 CORS 中间件
    app.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    # 静态文件和模板
    app.mount("/static", StaticFiles(directory="frontend/static"), name="static")
    templates = Jinja2Templates(directory="frontend")

    # 常量
    OUTPUT_DIR = os.getenv("OUTPUT_DIR", "./outputs")
    DOCS_DIR = os.getenv("DOCS_DIR", "./docs")

    # 启动事件
    @app.on_event("startup")
    async def startup_event():
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        app.mount("/outputs", StaticFiles(directory=OUTPUT_DIR), name="outputs")
        os.makedirs(DOCS_DIR, exist_ok=True)
        logger.info("YCM Researcher Enhanced 启动完成")

    # 路由
    @app.get("/", response_class=HTMLResponse)
    async def read_root(request: Request):
        return templates.TemplateResponse("index.html", {"request": request})

    # 研究请求模型
    class ResearchRequest(BaseModel):
        query: str
        report_type: str = ReportType.ResearchReport.value
        report_source: str = ReportSource.Web.value
        tone: str = Tone.Objective.value
        agent: str = "researcher"
        query_domains: List[str] = []

    # 研究路由
    @app.post("/api/research")
    async def conduct_research(request: ResearchRequest):
        """
        执行研究任务并返回结果

        Args:
            request: 研究请求

        Returns:
            研究结果
        """
        try:
            # 生成唯一的研究 ID
            research_id = f"{request.query.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

            # 创建研究代理
            researcher = YCMResearcher(
                query=request.query,
                report_type=request.report_type,
                report_source=request.report_source,
                tone=request.tone,
                log_handler=progress_handler,
                query_domains=request.query_domains,
            )

            # 发送开始研究的消息
            await progress_handler.send_progress(f"开始研究: {request.query}")

            # 执行研究
            research_result = await researcher.conduct_research()

            # 发送研究完成的消息
            await progress_handler.send_progress("研究完成，正在生成报告...")

            # 生成报告
            report = await researcher.write_report()

            # 保存研究结果到文件
            output_file = f"cache/research/{research_id}.json"

            # 收集研究过程中的日志
            research_logs = []
            log_file = "logs/research.log"
            if os.path.exists(log_file):
                try:
                    with open(log_file, "r", encoding="utf-8") as log_f:
                        for line in log_f:
                            if research_id in line or request.query in line:
                                research_logs.append(line.strip())
                except Exception as log_err:
                    logger.error(f"读取研究日志时出错: {str(log_err)}")

            # 将报告格式化为结构化的 JSON 对象
            structured_report = {
                "title": f"研究报告: {request.query}",
                "query": request.query,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "content": report,
                "sources": list(researcher.visited_urls),
                "research_logs": research_logs,  # 添加研究日志
            }

            with open(output_file, "w", encoding="utf-8") as f:
                # 确保 JSON 序列化时使用正确的编码
                json.dump(structured_report, f, ensure_ascii=False, indent=2)

            # 发送报告生成完成的消息
            await progress_handler.send_progress("报告生成完成")

            # 发送研究完成的消息，包含完整报告
            await progress_handler.websocket_manager.broadcast(
                {
                    "type": "research_complete",
                    "data": {
                        "status": "success",
                        "research_id": research_id,
                        "report": structured_report,
                    },
                }
            )

            return {
                "status": "success",
                "research_id": research_id,
                "report": structured_report,
            }
        except Exception as e:
            logger.error(f"研究过程中出错: {str(e)}", exc_info=True)
            # 发送错误消息
            await progress_handler.send_progress(f"研究过程中出错: {str(e)}")
            return {"status": "error", "message": f"研究过程中出错: {str(e)}"}

    # 获取研究报告路由
    @app.get("/api/research/{research_id}")
    async def get_research(research_id: str):
        """
        获取指定 ID 的研究报告

        Args:
            research_id: 研究 ID

        Returns:
            研究报告
        """
        try:
            file_path = f"cache/research/{research_id}.json"
            if not os.path.exists(file_path):
                raise HTTPException(status_code=404, detail="研究报告不存在")

            with open(file_path, "r", encoding="utf-8") as f:
                report = json.load(f)

            # 處理新舊格式的報告
            if isinstance(report, str):
                # 舊格式：報告是純文本
                structured_report = {
                    "title": f"研究報告",
                    "query": research_id.split("_")[0].replace("_", " "),
                    "timestamp": datetime.fromtimestamp(
                        os.path.getmtime(file_path)
                    ).strftime("%Y-%m-%d %H:%M:%S"),
                    "content": report,
                    "sources": [],
                }
                report = structured_report

            return {"status": "success", "report": report}
        except Exception as e:
            logger.error(f"获取研究报告时出错: {str(e)}", exc_info=True)
            # 发送错误消息
            await progress_handler.send_progress(f"获取研究报告时出错: {str(e)}")
            return {"status": "error", "message": f"获取研究报告时出错: {str(e)}"}

    # 获取所有保存的研究报告摘要路由
    @app.get("/api/research")
    async def get_saved_research_summaries():
        """
        获取所有保存的研究报告的摘要列表

        Returns:
            研究报告摘要列表
        """
        try:
            research_dir = Path("cache/research")
            if not research_dir.exists():
                return {"status": "success", "summaries": []}

            summaries = []
            for file_path in research_dir.glob("*.json"):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        report = json.load(f)

                    research_id = file_path.stem

                    # 兼容新旧格式的报告
                    if isinstance(report, str):
                        # 旧格式：报告是纯文本
                        title = "未知标题"
                        query = "未知查询"
                    else:
                        # 新格式：报告是结构化的 JSON 对象
                        title = report.get("title", "未知标题")
                        query = report.get("query", "未知查询")

                    summary = {
                        "id": research_id,
                        "title": title,
                        "query": query,
                        "date": datetime.fromtimestamp(
                            file_path.stat().st_mtime
                        ).strftime("%Y-%m-%d %H:%M:%S"),
                    }
                    summaries.append(summary)
                except Exception as e:
                    logger.error(f"处理研究报告 {file_path} 时出错: {str(e)}")

            # 按日期降序排序
            summaries.sort(key=lambda x: x["date"], reverse=True)

            return {"status": "success", "summaries": summaries}
        except Exception as e:
            logger.error(f"获取研究报告摘要时出错: {str(e)}", exc_info=True)
            # 发送错误消息
            await progress_handler.send_progress(f"获取研究报告摘要时出错: {str(e)}")
            return {"status": "error", "message": f"获取研究报告摘要时出错: {str(e)}"}

    # API 测试路由
    @app.get("/api/test")
    async def test_api():
        return {"status": "success", "message": "API 工作正常"}

    # WebSocket 端点
    @app.websocket("/ws")
    async def websocket_endpoint(websocket: WebSocket):
        await manager.connect(websocket)
        try:
            while True:
                data = await websocket.receive_text()
                # 这里可以处理从客户端接收的消息
        except WebSocketDisconnect:
            manager.disconnect(websocket)

    # 下载研究报告路由
    @app.get("/api/research/{research_id}/download")
    async def download_research(research_id: str, format: str = "markdown"):
        """
        下载指定格式的研究报告

        Args:
            research_id: 研究 ID
            format: 下载格式 (markdown, pdf, docx, json)

        Returns:
            文件下载响应
        """
        try:
            file_path = f"cache/research/{research_id}.json"
            if not os.path.exists(file_path):
                raise HTTPException(status_code=404, detail="研究报告不存在")

            with open(file_path, "r", encoding="utf-8") as f:
                report = json.load(f)

            # 處理新舊格式的報告
            if isinstance(report, str):
                # 舊格式：報告是純文本
                title = "研究報告"
                content = report
            else:
                # 新格式：報告是結構化的 JSON 對象
                title = report.get("title", "研究報告")
                content = report.get("content", "")

            # 根据请求的格式生成对应的文件
            if format == "json":
                # 直接返回 JSON 文件
                return FileResponse(
                    file_path,
                    media_type="application/json",
                    filename=f"{title.replace(' ', '_')}.json",
                )
            elif format == "markdown":
                # 生成 Markdown 文件
                md_file = f"cache/research/{research_id}.md"
                with open(md_file, "w", encoding="utf-8") as f:
                    f.write(content)

                return FileResponse(
                    md_file,
                    media_type="text/markdown",
                    filename=f"{title.replace(' ', '_')}.md",
                )
            elif format == "pdf":
                # 生成 PDF 文件
                try:
                    import pdfkit
                    from weasyprint import HTML, CSS

                    # 转换 Markdown 为 HTML
                    html_content = markdown.markdown(
                        content, extensions=["tables", "fenced_code"]
                    )

                    # 添加基本样式
                    styled_html = f"""
                    <!DOCTYPE html>
                    <html>
                    <head>
                        <meta charset="UTF-8">
                        <title>{title}</title>
                        <style>
                            body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                            h1, h2, h3, h4, h5, h6 {{ color: #333; margin-top: 20px; }}
                            h1 {{ font-size: 24px; }}
                            h2 {{ font-size: 20px; }}
                            h3 {{ font-size: 18px; }}
                            p {{ margin-bottom: 10px; }}
                            code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 4px; }}
                            pre {{ background-color: #f5f5f5; padding: 10px; border-radius: 4px; overflow-x: auto; }}
                            blockquote {{ border-left: 4px solid #ddd; padding-left: 10px; color: #666; }}
                            table {{ border-collapse: collapse; width: 100%; }}
                            th, td {{ border: 1px solid #ddd; padding: 8px; }}
                            th {{ background-color: #f2f2f2; }}
                            img {{ max-width: 100%; }}
                            a {{ color: #0066cc; text-decoration: none; }}
                        </style>
                    </head>
                    <body>
                        <h1>{title}</h1>
                        {html_content}
                    </body>
                    </html>
                    """

                    # 保存 HTML 文件
                    html_file = f"cache/research/{research_id}.html"
                    with open(html_file, "w", encoding="utf-8") as f:
                        f.write(styled_html)

                    # 生成 PDF 文件
                    pdf_file = f"cache/research/{research_id}.pdf"

                    # 尝试使用 WeasyPrint 生成 PDF
                    try:
                        HTML(string=styled_html).write_pdf(pdf_file)
                    except Exception as e:
                        logger.error(f"使用 WeasyPrint 生成 PDF 失败: {str(e)}")
                        # 如果 WeasyPrint 失败，尝试使用 pdfkit
                        try:
                            pdfkit.from_string(styled_html, pdf_file)
                        except Exception as e2:
                            logger.error(f"使用 pdfkit 生成 PDF 失败: {str(e2)}")
                            raise HTTPException(status_code=500, detail="生成 PDF 失败")

                    return FileResponse(
                        pdf_file,
                        media_type="application/pdf",
                        filename=f"{title.replace(' ', '_')}.pdf",
                    )
                except ImportError as e:
                    logger.error(f"缺少生成 PDF 所需的库: {str(e)}")
                    raise HTTPException(
                        status_code=500, detail="服务器缺少生成 PDF 所需的库"
                    )
            elif format == "docx":
                # 生成 Word 文件
                try:
                    from docx import Document
                    from docx.shared import Pt, Inches

                    # 创建 Word 文档
                    doc = Document()

                    # 设置文档标题
                    doc.add_heading(title, level=1)

                    # 将 Markdown 内容转换为 Word 格式
                    # 简单处理：按行分割，根据 Markdown 语法添加到文档
                    lines = content.split("\n")
                    for line in lines:
                        # 处理标题
                        if line.startswith("# "):
                            doc.add_heading(line[2:], level=1)
                        elif line.startswith("## "):
                            doc.add_heading(line[3:], level=2)
                        elif line.startswith("### "):
                            doc.add_heading(line[4:], level=3)
                        elif line.startswith("#### "):
                            doc.add_heading(line[5:], level=4)
                        elif line.startswith("##### "):
                            doc.add_heading(line[6:], level=5)
                        elif line.startswith("###### "):
                            doc.add_heading(line[7:], level=6)
                        # 处理列表项
                        elif line.startswith("- ") or line.startswith("* "):
                            doc.add_paragraph(line[2:], style="List Bullet")
                        elif re.match(r"^\d+\. ", line):
                            doc.add_paragraph(
                                re.sub(r"^\d+\. ", "", line), style="List Number"
                            )
                        # 处理引用
                        elif line.startswith("> "):
                            p = doc.add_paragraph(line[2:])
                            p.style = "Quote"
                        # 处理普通段落
                        elif line.strip() != "":
                            doc.add_paragraph(line)
                        # 处理空行
                        else:
                            doc.add_paragraph()

                    # 保存 Word 文件
                    docx_file = f"cache/research/{research_id}.docx"
                    doc.save(docx_file)

                    return FileResponse(
                        docx_file,
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        filename=f"{title.replace(' ', '_')}.docx",
                    )
                except ImportError as e:
                    logger.error(f"缺少生成 Word 文件所需的库: {str(e)}")
                    raise HTTPException(
                        status_code=500, detail="服务器缺少生成 Word 文件所需的库"
                    )
            else:
                raise HTTPException(status_code=400, detail=f"不支持的格式: {format}")
        except Exception as e:
            logger.error(f"下载研究报告时出错: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"下载研究报告时出错: {str(e)}")

    # 主函数
    if __name__ == "__main__":
        logger.info("启动 YCM Researcher Enhanced 服务器...")
        uvicorn.run(app, host="localhost", port=8080)
except Exception as e:
    logger.error(f"导入错误: {str(e)}", exc_info=True)
    print(f"导入错误: {str(e)}")
    print("请确保已安装所有必要的依赖项: pip install fastapi uvicorn")
